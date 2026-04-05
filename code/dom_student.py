"""
=============================================================================
dom_student.py  —  Discrete Ordinates Method (DOM) for 1-D Leaf Canopy
                   Radiative Transfer (F)
Authors         -  Claude And I (Ranga B. Myneni)
=============================================================================

PURPOSE
-------
This code solves the one-dimensional (1-D) monochromatic radiative transfer
equation (RTE) for a horizontally-uniform, infinitely-extended leaf canopy
using the Discrete Ordinates Method (DOM).

It produces:
  1.  Energy balance tables (uncollided + collided) written to a Word document.
  2.  Optionally, a table of BRF or HDRF values at a user-defined angular
      grid, written to a second Word document.  The interpolation is carried
      out using the DOM method described in the companion document.

The code does NOT produce any figures.  Run it, supply the inputs at the
bottom of the file (Section 0), and two .docx files will appear in the
working directory.

REQUIREMENTS
------------
  Python  >= 3.9
  numpy   >= 1.22
  scipy   >= 1.8   (only if you choose the optional spline cross-check)
  python-docx >= 1.1   (pip install python-docx)

STRUCTURE
---------
  Section 0  — User inputs  (edit this section only)
  Section 1  — Gauss-Legendre quadrature
  Section 2  — Leaf extinction coefficient  G
  Section 3  — Bi-Lambertian phase function  Gamma
  Section 4  — Uncollided field  (Beer-Lambert sweeps)
  Section 5  — First Collision Source  Q = Q1 + Q2 + Q3
  Section 6  — Collided field  (Diamond-Difference + iterative scattering)
  Section 7  — Energy balance
  Section 8  — DOM interpolation into arbitrary view direction
  Section 9  — Write energy-balance Word document
  Section 10 — Write BRF / HDRF Word document
  Section 11 — Master solve routine
  Section 12 — Main entry point
=============================================================================
"""

# ── Standard library ──────────────────────────────────────────────────────────
import sys
import math
import time

# ── Third-party ───────────────────────────────────────────────────────────────
import numpy as np

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit(
        "ERROR: python-docx is not installed.\n"
        "Install it with:  pip install python-docx\n"
    )


# =============================================================================
# SECTION 0 — USER INPUTS   <<<  EDIT THIS SECTION ONLY  >>>
# =============================================================================

# ── Spectral band parameters ──────────────────────────────────────────────────
#   rho_L   : leaf single-sided reflectance   (dimensionless, 0–1)
#   tau_L   : leaf single-sided transmittance  (dimensionless, 0–1)
#   rho_g   : Lambertian ground reflectance    (dimensionless, 0–1)
#   f_dir   : fraction of downwelling irradiance that is direct solar beam
#             f_dir = 0 → all diffuse sky; f_dir = 1 → all direct solar beam
#             If f_dir < 1, the output is labelled HDRF; if f_dir = 1, BRF.
#
#   One dict per band.  Add or remove bands freely.

SPECTRAL_BANDS = {
    "RED": {
        "rho_L":  0.06,    # leaf reflectance (RED)
        "tau_L":  0.04,    # leaf transmittance (RED)
        "rho_g":  0.10,    # ground reflectance (RED)
        "f_dir":  0.70,    # direct fraction (RED)
    },
    "NIR": {
        "rho_L":  0.525,   # leaf reflectance (NIR)
        "tau_L":  0.450,   # leaf transmittance (NIR)
        "rho_g":  0.20,    # ground reflectance (NIR)
        "f_dir":  0.70,    # direct fraction (NIR)
    },
}

# ── Canopy parameters ─────────────────────────────────────────────────────────
LAI_VALUES = [1.5, 4.0]   # canopy Leaf Area Index values to compute (list)

# ── Illumination geometry ─────────────────────────────────────────────────────
#   THETA_SUN_DEG : solar polar angle in degrees, measured from the UPWARD
#                  zenith.  The sun is always below the horizon from above, so
#                  THETA_SUN_DEG must be in (90, 180].
#                    180° → sun directly overhead (nadir incidence from above)
#                    140° → sun 40° from nadir (standard test case)
#                     91° → sun almost at the horizon
THETA_SUN_DEG = 140.0   # degrees from upward zenith
PHI_SUN_DEG   =   0.0   # solar azimuth (degrees).  0 = principal plane.

# ── Incoming irradiance ───────────────────────────────────────────────────────
F_IN = 1.0   # total downwelling irradiance at the top of the canopy [W m-2]
             # (results are normalised by F_IN, so the actual value does not
             # change any reflectance or transmittance — only absolute fluxes)

# ── Numerical resolution ──────────────────────────────────────────────────────
N_MU  = 16   # number of Gauss-Legendre polar directions on the full sphere
             # (must be even; N_MU/2 downward + N_MU/2 upward)
N_PHI = 16   # number of uniform azimuthal directions
N_LAY = 50   # number of canopy layers (cells)

# ── Solver convergence ────────────────────────────────────────────────────────
TOL      = 0.01   # convergence tolerance (1 % relative change in source)
MAX_ITER = 500    # safety limit on number of scattering iterations

# ── Output file names ─────────────────────────────────────────────────────────
ENERGY_BALANCE_DOCX = "energy_balance.docx"
BRF_HDRF_DOCX       = "brf_hdrf_table.docx"

# ── Interpolation grid ────────────────────────────────────────────────────────
#   The BRF / HDRF table is evaluated at these view-direction grid points.
#   theta_v : view polar angle from nadir  (degrees, 0 = nadir, 90 = horizon)
#   phi_v   : view azimuth angle           (degrees, 0–360)
THETA_V_STEP = 2    # polar step in degrees
PHI_V_STEP   = 4    # azimuth step in degrees


# =============================================================================
# SECTION 1 — GAUSS-LEGENDRE QUADRATURE
# =============================================================================

def setup_quadrature(N, M):
    """
    Build the angular quadrature for the discrete ordinates solver.

    POLAR DIRECTIONS  (Gauss-Legendre)
    ------------------------------------
    We need N quadrature points for the full sphere in the polar cosine
    variable mu = cos(theta), where theta is measured from the UPWARD zenith:

        mu in [-1, 0)  →  downward-travelling directions   (sun, sky, diffuse)
        mu in  (0, 1]  →  upward-travelling directions     (reflected radiance)

    numpy.polynomial.legendre.leggauss(N) returns N nodes on (-1,1) and the
    corresponding weights.  The nodes come out in ascending order, so:
        index 1  ..  N/2  :  downward (mu < 0), ordered from most-downward to nadir
        index N/2+1 ..  N :  upward   (mu > 0), ordered from nadir to most-upward

    We use 1-based indexing throughout.  Index 0 is a dummy (never used).

    AZIMUTHAL DIRECTIONS  (uniform)
    --------------------------------
    M uniform azimuth angles phi_m = (m - 1) * 2*pi/M for m = 1 .. M.
    All azimuthal weights are equal:  w_phi = 2*pi/M.

    PARAMETERS
    ----------
    N : int   number of polar quadrature points on the FULL sphere (must be even)
    M : int   number of azimuthal quadrature points

    RETURNS
    -------
    mu    : ndarray (N+1,)   polar cosines,  index 1..N are valid  (index 0 = 0)
    phi   : ndarray (M+1,)   azimuth angles (radians), index 1..M valid
    w_mu  : ndarray (N+1,)   Gauss-Legendre weights for polar integration
    w_phi : float            uniform weight for one azimuthal strip  (= 2*pi/M)
    """
    # Gauss-Legendre nodes and weights on (-1, 1)
    nodes, weights = np.polynomial.legendre.leggauss(N)   # ascending order

    # Allocate 1-based arrays
    mu   = np.zeros(N + 1)
    w_mu = np.zeros(N + 1)

    mu[1:N+1]   = nodes    # index 1..N
    w_mu[1:N+1] = weights

    # Uniform azimuth grid
    phi   = np.zeros(M + 1)
    phi[1:M+1] = (np.arange(M) / M) * 2.0 * np.pi   # 0, 2pi/M, 4pi/M, ...
    w_phi = 2.0 * np.pi / M

    # --- Sanity checks ---
    assert abs(np.sum(w_mu[1:N+1]) - 2.0) < 1e-10, "Polar weights must sum to 2"
    half = N // 2
    flux_check = np.sum(np.abs(mu[1:half+1]) * w_mu[1:half+1])
    # Note: this check only converges for large N; allow looser tolerance
    assert abs(flux_check - 0.5) < 0.05, "Upward flux weight should be near 0.5"

    w_phi = w_phi_scalar(M)
    return mu, phi, w_mu, w_phi


def w_phi_scalar(M):
    """Return the scalar azimuthal weight  2*pi / M."""
    return 2.0 * np.pi / M


# =============================================================================
# SECTION 2 — LEAF EXTINCTION COEFFICIENT  G
# =============================================================================

def G_extinction():
    """
    Extinction coefficient G for a uniform leaf-angle distribution (spherical).

    For a uniform leaf-angle distribution (all leaf-normals equally probable),
    the extinction coefficient is exactly  G = 0.5  for ALL directions.

    This is the standard assumption used throughout this code.

    RETURNS
    -------
    float  : 0.5
    """
    return 0.5


# =============================================================================
# SECTION 3 — BI-LAMBERTIAN PHASE FUNCTION  Gamma
# =============================================================================

def gamma_scalar(mu_i, phi_i, mu_j, phi_j, omega_L, tau_L):
    """
    Bi-Lambertian scattering phase function.

    The bi-Lambertian phase function describes how a single leaf element
    scatters radiation from an incident direction (mu_i, phi_i) into a
    scattered direction (mu_j, phi_j).

    DEFINITION
    ----------
    The scattering angle beta between the two directions satisfies:

        cos(beta) = mu_i * mu_j
                  + sqrt(1 - mu_i^2) * sqrt(1 - mu_j^2) * cos(phi_i - phi_j)

    The phase function is then:

        Gamma(i -> j) = (omega_L / (3*pi)) * [sin(beta) - beta*cos(beta)]
                      + (tau_L  / 3)       * cos(beta)

    where:
        omega_L = rho_L + tau_L   (single-scattering albedo)
        tau_L                     (leaf transmittance, determines forward peak)

    PARAMETERS
    ----------
    mu_i, phi_i : float   incident direction (polar cosine, azimuth in radians)
    mu_j, phi_j : float   scattered direction
    omega_L     : float   leaf single-scattering albedo = rho_L + tau_L
    tau_L       : float   leaf transmittance

    RETURNS
    -------
    float : Gamma value  [sr-1]
    """
    sin_i = math.sqrt(max(0.0, 1.0 - mu_i**2))
    sin_j = math.sqrt(max(0.0, 1.0 - mu_j**2))
    cos_beta = np.clip(
        mu_i * mu_j + sin_i * sin_j * math.cos(phi_i - phi_j),
        -1.0, 1.0
    )
    beta = math.acos(cos_beta)
    return (omega_L / (3.0 * math.pi)) * (math.sin(beta) - beta * cos_beta) \
           + (tau_L / 3.0) * cos_beta


def precompute_gamma_quad(mu, phi, N, M, omega_L, tau_L):
    """
    Pre-compute the full Gauss-to-Gauss phase-function table.

    Gamma[n, m, i, j]  =  Gamma( (mu_n, phi_m) -> (mu_i, phi_j) )

    for all pairs of Gauss directions.  This is the most expensive part of
    the setup; it is called once per band and cached.

    PARAMETERS
    ----------
    mu     : ndarray (N+1,)   polar cosines (1-based)
    phi    : ndarray (M+1,)   azimuths in radians (1-based)
    N, M   : int              quadrature sizes
    omega_L, tau_L : float

    RETURNS
    -------
    G_qq : ndarray (N+1, M+1, N+1, M+1)
           Gauss-to-Gauss phase function table
    """
    print(f"  Pre-computing Gamma (N={N}, M={M}: {N*M*N*M} evaluations) ...")
    G_qq = np.zeros((N+1, M+1, N+1, M+1))
    for n in range(1, N+1):
        for m in range(1, M+1):
            for i in range(1, N+1):
                for j in range(1, M+1):
                    G_qq[n, m, i, j] = gamma_scalar(
                        mu[n], phi[m], mu[i], phi[j], omega_L, tau_L)
    print("  Done.\n")
    return G_qq


def precompute_gamma_solar(solar_mu, solar_phi, mu, phi, N, M, omega_L, tau_L):
    """
    Pre-compute the solar-to-Gauss phase-function column.

    G_sol[i, j]  =  Gamma( (solar_mu, solar_phi) -> (mu_i, phi_j) )

    RETURNS
    -------
    G_sol : ndarray (N+1, M+1)
    """
    G_sol = np.zeros((N+1, M+1))
    for i in range(1, N+1):
        for j in range(1, M+1):
            G_sol[i, j] = gamma_scalar(
                solar_mu, solar_phi, mu[i], phi[j], omega_L, tau_L)
    return G_sol


# =============================================================================
# SECTION 4 — UNCOLLIDED FIELD
# =============================================================================

def solve_uncollided(mu, phi, w_mu, w_phi, N, M, K,
                     solar_mu, solar_phi, f_dir, F_in, rho_g, G, dL):
    """
    Solve for the uncollided intensity field.

    The uncollided field consists of two contributions:
      (a) Direct solar beam    I0_dir(L)  — scalar, only along solar direction
      (b) Diffuse sky + ground I0[n,m,L]  — defined at all Gauss directions

    Both follow Beer-Lambert attenuation:

        I0_dir(L)      =  I_beam * exp( -G * L / |mu_solar| )
        I0[n,m,L]      : downward sweep + ground reflection + upward sweep

    LAYER INDEXING
    --------------
    L = 0       : canopy top
    L = LAI     : canopy bottom (ground)

    We use K+2 edge points: k = 1 (top) to k = K+1 (bottom).
    Cell centres are at k + 1/2 for k = 1 .. K.

    PARAMETERS
    ----------
    mu, phi, w_mu, w_phi : quadrature arrays
    N, M, K  : integers     quadrature and layer counts
    solar_mu, solar_phi : float   solar direction cosines
    f_dir    : float   direct fraction of incoming irradiance
    F_in     : float   total downwelling irradiance at top
    rho_g    : float   Lambertian ground reflectance
    G        : float   extinction coefficient (= 0.5)
    dL       : float   layer thickness = LAI / K

    RETURNS
    -------
    I0_dir : ndarray (K+2,)        direct-beam scalar intensity at cell edges
    I0     : ndarray (N+1,M+1,K+2) diffuse uncollided intensity at cell edges
    Fd_dir : ndarray (K+2,)        direct-beam downward irradiance
    Fd_dif : ndarray (K+2,)        diffuse downward irradiance (uncollided)
    Fu_unc : ndarray (K+2,)        upward irradiance (uncollided, from ground)
    """
    half = N // 2   # N/2 downward + N/2 upward directions

    # ── 4a. Direct solar beam ─────────────────────────────────────────────────
    #
    # The direct beam travels along the solar direction (solar_mu < 0, downward).
    # Its intensity just above the canopy is:
    #
    #   I_beam = f_dir * F_in / |solar_mu|
    #
    # Meaning: the direct-beam irradiance f_dir*F_in arrives along one direction,
    # so to convert irradiance to intensity we divide by |solar_mu|.
    #
    # Beer-Lambert attenuation is ANALYTICAL for the uncollided field because
    # there is no scattering source term.  No sweep is needed — just evaluate
    # the exponential directly at each edge k:
    #
    #   I0_dir(k) = I_beam * exp( -G * (k-1)*dL / |solar_mu| )
    #
    # (k-1)*dL is the cumulative LAI from the canopy top (k=1) to edge k.
    #
    abs_mu_sol = abs(solar_mu)
    I_beam = f_dir * F_in / abs_mu_sol if abs_mu_sol > 1e-10 else 0.0

    # Step 3a (C4-P3 Slide 12): Sweep down for direct solar
    # I0dir(Lk=1,  Omega_o) = I0dir(L=0,  Omega_o)  <- upper BC
    # I0dir(Lk=2,  Omega_o) = I0dir(Lk=1, Omega_o) * exp[-(1/|mu_o|)*G(Omega_o)*dL]
    # I0dir(Lk=3,  Omega_o) = I0dir(Lk=2, Omega_o) * exp[-(1/|mu_o|)*G(Omega_o)*dL]
    # ...
    # I0dir(Lk=K+1,Omega_o) = I0dir(Lk=K, Omega_o) * exp[-(1/|mu_o|)*G(Omega_o)*dL]
    exp_step_sol = math.exp(-G * dL / abs_mu_sol)   # attenuation per layer
    I0_dir = np.zeros(K + 2)
    I0_dir[1] = I_beam                               # upper boundary condition
    for k in range(1, K + 1):
        I0_dir[k+1] = I0_dir[k] * exp_step_sol      # recurrence sweep downward

    # Direct downward irradiance:  F_dir(k) = |mu_o| * I0_dir(k)
    Fd_dir = np.zeros(K + 2)
    for k in range(1, K + 2):
        Fd_dir[k] = abs_mu_sol * I0_dir[k]

    # ── 4b. Diffuse sky — downward Beer-Lambert (no sweep needed) ────────────
    #
    # The diffuse sky illuminates the canopy isotropically.  Its radiance at
    # the canopy top is:
    #
    #   I_sky = (1 - f_dir) * F_in / pi   [W m-2 sr-1]
    #
    # For each downward Gauss direction (n = 1..N/2), the uncollided intensity
    # at each edge k is obtained analytically:
    #
    #   I0[n, m, k] = I_sky * exp( -G * (k-1)*dL / |mu_n| )
    #
    I0 = np.zeros((N + 1, M + 1, K + 2))
    I_sky = (1.0 - f_dir) * F_in / math.pi

    # Step 3b (C4-P3 Slide 13): Sweep down for each of the (N/2)*M downward
    # Gauss directions Omega_ij:
    # I0dif(Lk=1,  Omega_ij) = I0dif(L=0,  Omega_ij)  <- upper BC = I_sky
    # I0dif(Lk=2,  Omega_ij) = I0dif(Lk=1, Omega_ij) * exp[-(1/|mu_i|)*G(Omega_ij)*dL]
    # ...and so on till k=K+1
    for n in range(1, half + 1):       # downward: mu_n < 0
        abs_mu_n = abs(mu[n])
        exp_step_n = math.exp(-G * dL / abs_mu_n)   # attenuation per layer
        for m in range(1, M + 1):
            I0[n, m, 1] = I_sky                      # upper boundary condition
            for k in range(1, K + 1):
                I0[n, m, k+1] = I0[n, m, k] * exp_step_n   # recurrence

    # ── 4c. Downward irradiance from diffuse sky ──────────────────────────────
    #
    # Integrate over the N/2 downward Gauss directions:
    #   Fd_dif(k) = SUM_{n=1}^{N/2} SUM_{m=1}^{M}  w_n * w_phi * |mu_n| * I0[n,m,k]
    #
    Fd_dif = np.zeros(K + 2)
    for k in range(1, K + 2):
        for n in range(1, half + 1):
            for m in range(1, M + 1):
                Fd_dif[k] += w_mu[n] * w_phi * abs(mu[n]) * I0[n, m, k]

    # ── 4d. Ground boundary condition (Lambertian reflection) ─────────────────
    #
    # Total downward flux at the ground:
    F_down_gnd = Fd_dif[K+1] + Fd_dir[K+1]
    #
    # Lambertian reflection: the ground reflects uniformly in all upward
    # directions with the same radiance:
    #
    #   I0_up_gnd = (rho_g / pi) * F_down_gnd
    #
    I0_up_gnd = (rho_g / math.pi) * F_down_gnd

    # ── 4e. Upward Beer-Lambert from ground (no sweep needed) ─────────────────
    #
    # The ground-reflected radiance attenuates upward through the canopy.
    # For each upward Gauss direction (n = N/2+1..N), the intensity at edge k
    # (measured upward from the ground at k=K+1) is:
    #
    #   I0[n, m, k] = I0_up_gnd * exp( -G * (K+1-k)*dL / |mu_n| )
    #
    # (K+1-k)*dL = distance from edge k up to the ground at edge K+1.
    #
    # Step 5 (C4-P3 Slide 16): Sweep up for each of the (N/2)*M upward
    # Gauss directions Omega_ij:
    # I0(Lk=K+1, Omega_ij) = known from lower BC (Step 4)
    # I0(Lk=K,   Omega_ij) = I0(Lk=K+1, Omega_ij) * exp[-(1/|mu_i|)*G(Omega_ij)*dL]
    # I0(Lk=K-1, Omega_ij) = I0(Lk=K,   Omega_ij) * exp[-(1/|mu_i|)*G(Omega_ij)*dL]
    # ...and so on till k=1
    Fu_unc = np.zeros(K + 2)
    for n in range(half + 1, N + 1):   # upward: mu_n > 0
        abs_mu_n = abs(mu[n])
        exp_step_n = math.exp(-G * dL / abs_mu_n)   # attenuation per layer
        for m in range(1, M + 1):
            I0[n, m, K+1] = I0_up_gnd               # lower boundary condition
            for k in range(K, 0, -1):
                I0[n, m, k] = I0[n, m, k+1] * exp_step_n   # recurrence

    # ── 4f. Upward irradiance from ground-reflected uncollided field ──────────
    #
    #   Fu_unc(k) = SUM_{n=N/2+1}^{N} SUM_{m=1}^{M}  w_n * w_phi * |mu_n| * I0[n,m,k]
    #
    for k in range(1, K + 2):
        for n in range(half + 1, N + 1):
            for m in range(1, M + 1):
                Fu_unc[k] += w_mu[n] * w_phi * abs(mu[n]) * I0[n, m, k]

    return I0_dir, I0, Fd_dir, Fd_dif, Fu_unc


# =============================================================================
# SECTION 5 — FIRST COLLISION SOURCE  Q = Q1 + Q2 + Q3
# =============================================================================

def compute_first_collision_source(mu, phi, w_mu, w_phi, N, M, K,
                                    solar_mu, solar_phi, I0_dir, I0,
                                    G_sol, G_qq, omega_L):
    """
    Compute the First Collision Source (FCS) Q at every cell centre.

    When a photon undergoes its FIRST scattering event inside the canopy, it
    contributes to the collided field.  The source has three terms:

    Q1 : direct solar beam scattered into each Gauss direction
    Q2 : downward diffuse sky scattered into each Gauss direction
    Q3 : upward ground-reflected uncollided flux scattered into each Gauss direction

    Each term has the form:

        Q(Lk+1/2, Omega_ij) = (1/pi) * Gamma(source -> Omega_ij) * I_source(Lk+1/2)

    For Q1 (direct solar):
        Q1[i,j,k] = (1/pi) * Gamma(solar -> Omega_ij) * I0_dir(Lk+1/2)

    For Q2 (diffuse downward, n=1..N/2):
        Q2[i,j,k] = (1/pi) * SUM_{n=1}^{N/2}  SUM_{m=1}^{M}
                              w_n * w_phi * Gamma(nm -> ij) * I0_dif[n,m,k]

    For Q3 (upward uncollided, n=N/2+1..N):
        Q3[i,j,k] = (1/pi) * SUM_{n=N/2+1}^{N}  SUM_{m=1}^{M}
                              w_n * w_phi * Gamma(nm -> ij) * I0[n,m,k]

    Cell-centre values are obtained as the arithmetic mean of the two
    surrounding edge values:
        I0_dir_c[k] = (I0_dir[k] + I0_dir[k+1]) / 2  for k = 1..K

    PARAMETERS
    ----------
    mu, phi, w_mu, w_phi : quadrature
    N, M, K  : integers
    solar_mu, solar_phi : float
    I0_dir   : ndarray (K+2,)         direct-beam scalar intensity (edges)
    I0       : ndarray (N+1,M+1,K+2)  diffuse uncollided intensity (edges)
    G_sol    : ndarray (N+1,M+1)      Gamma(solar -> Gauss)
    G_qq     : ndarray (N+1,M+1,N+1,M+1)  Gamma(Gauss -> Gauss)
    omega_L  : float   single-scattering albedo

    RETURNS
    -------
    Q : ndarray (N+1, M+1, K)   FCS at cell centres k=1..K
    """
    half = N // 2
    Q = np.zeros((N + 1, M + 1, K))

    # Cell-centre intensities (average of surrounding edges)
    I0_dir_c = 0.5 * (I0_dir[1:K+1] + I0_dir[2:K+2])             # (K,)
    I0_c     = 0.5 * (I0[:, :, 1:K+1] + I0[:, :, 2:K+2])         # (N+1,M+1,K)

    for i in range(1, N + 1):
        for j in range(1, M + 1):
            # Q1: direct solar beam
            q1 = (1.0 / math.pi) * G_sol[i, j] * I0_dir_c        # (K,)

            # Q2: downward diffuse sky (n = 1 .. N/2)
            q2 = np.zeros(K)
            for n in range(1, half + 1):
                for m in range(1, M + 1):
                    q2 += w_mu[n] * w_phi * G_qq[n, m, i, j] * I0_c[n, m, :]
            q2 /= math.pi

            # Q3: upward uncollided (n = N/2+1 .. N)
            q3 = np.zeros(K)
            for n in range(half + 1, N + 1):
                for m in range(1, M + 1):
                    q3 += w_mu[n] * w_phi * G_qq[n, m, i, j] * I0_c[n, m, :]
            q3 /= math.pi

            Q[i, j, :] = q1 + q2 + q3

    return Q   # shape (N+1, M+1, K)


# =============================================================================
# SECTION 6 — COLLIDED FIELD  (Diamond-Difference + iterative scattering)
# =============================================================================

def solve_collided(mu, phi, w_mu, w_phi, N, M, K,
                   G, dL, rho_g, Q, G_qq, omega_L, tol, max_iter):
    """
    Solve for the collided intensity field by iterating on the multiple-
    scattering source S until convergence.

    ALGORITHM
    ---------
    1.  Initialise:  IC = 0,  S = 0
    2.  ITERATION LOOP:
        a. Compute total source  J[i,j,k] = Q[i,j,k] + S[i,j,k]
        b. Apply ground BC for collided field (Lambertian reflection of
           downward collided irradiance)
        c. Sweep all directions to obtain new IC at cell edges
        d. Compute new multiple-scattering source S from converged IC
        e. Check convergence  (relative change in S integrated over canopy)
    3.  Return converged IC

    DIAMOND-DIFFERENCE (DD) SCHEME
    --------------------------------
    For a downward direction (mu_n < 0), sweeping from k=1 (top) to k=K+1:

        IC[n,m,k+1] = (c * IC[n,m,k] + dd * J[n,m,k]) / 1
        where:
            half_f = 0.5 * G * dL / |mu_n|
            c      = (1 - half_f) / (1 + half_f)
            dd     = (dL / |mu_n|)  / (1 + half_f)

    For an upward direction (mu_n > 0), sweeping from k=K+1 (bottom) to k=1:

        IC[n,m,k] = c * IC[n,m,k+1] + dd * J[n,m,k]

    The cell-centre value is the mean of the two surrounding edges:
        IC_c[n,m,k] = (IC[n,m,k] + IC[n,m,k+1]) / 2

    GROUND BOUNDARY CONDITION (collided)
    ------------------------------------
    At the ground (k=K+1), the upward collided intensity is the Lambertian
    reflection of the total downward collided irradiance:

        IC[n,m,K+1] = (rho_g / pi) * F_col_down(K+1)

    where F_col_down = SUM over downward directions of w_n*w_phi*|mu_n|*IC[n,m,K+1]

    CONVERGENCE CRITERION
    ----------------------
    We check the relative change in the SCALAR IRRADIANCE (integral of S over
    all layers):

        bnd = |SUM(S_new) - SUM(S_old)| / max(|SUM(S_new)|, epsilon)

    Convergence when bnd < tol.

    PARAMETERS
    ----------
    (same conventions as Section 5)
    Q     : ndarray (N+1, M+1, K)   First Collision Source at cell centres
    G_qq  : ndarray (N+1,M+1,N+1,M+1)  Gauss-to-Gauss phase function
    tol   : float   convergence tolerance (relative)
    max_iter : int  safety limit

    RETURNS
    -------
    IC       : ndarray (N+1, M+1, K+2)   collided intensity at cell edges
    n_iter   : int                        number of iterations taken
    Fd_col   : ndarray (K+2,)             downward collided irradiance
    Fu_col   : ndarray (K+2,)             upward collided irradiance
    """
    half = N // 2
    IC   = np.zeros((N + 1, M + 1, K + 2))
    S    = np.zeros((N + 1, M + 1, K))        # multiple-scattering source

    # Pre-compute DD sweep coefficients for each direction (C4-P4, Slide 6).
    #
    # IMPORTANT SIGN CONVENTION (matching dom_rt.py and the slides exactly):
    # f is defined with the SIGNED mu_i, not |mu_i|:
    #
    #   f_ij = G(Omega_ij) * dL / mu_i
    #
    # For DOWNWARD directions (mu_i < 0): f_ij < 0
    #   a_ij = (1 + [1-alpha]*f_ij) / (1 - alpha*f_ij)   => a < 1 (attenuating)
    #   b_ij = f_ij / (G * (1 - alpha*f_ij))              => b < 0
    #
    #   Sweep equation: IC(Lk+1) = a_ij * IC(Lk) - b_ij * J(Lk+1/2)
    #                             = a*IC(k) + |b|*J     (source ADDS to IC) ✓
    #
    # For UPWARD directions (mu_i > 0): f_ij > 0
    #   c_ij = (1 - [1-alpha]*f_ij) / (1 + alpha*f_ij)   => c < 1 (attenuating)
    #   d_ij = f_ij / (G * (1 + alpha*f_ij))              => d > 0
    #
    #   Sweep equation: IC(Lk) = c_ij * IC(Lk+1) + d_ij * J(Lk+1/2)
    #                           (source ADDS to IC) ✓
    #
    # Using |mu_i| instead of mu_i for f would make f > 0 for downward,
    # giving a > 1 and b > 0, and the -b*J term would SUBTRACT the source
    # from IC — physically wrong and numerically unstable.
    #
    alpha = 0.5
    f_coef  = np.zeros(N + 1)
    a_coef  = np.zeros(N + 1)   # downward: IC(k+1) = a*IC(k) - b*J  (b<0, so +|b|*J)
    b_coef  = np.zeros(N + 1)
    c_coef  = np.zeros(N + 1)   # upward:   IC(k)   = c*IC(k+1) + d*J
    d_coef  = np.zeros(N + 1)
    for n in range(1, N + 1):
        f_coef[n] = G * dL / mu[n]   # SIGNED: negative for downward, positive for upward
        a_coef[n] = (1.0 + (1.0-alpha)*f_coef[n]) / (1.0 - alpha*f_coef[n])
        b_coef[n] = f_coef[n] / (G * (1.0 - alpha*f_coef[n]))
        c_coef[n] = (1.0 - (1.0-alpha)*f_coef[n]) / (1.0 + alpha*f_coef[n])
        d_coef[n] = f_coef[n] / (G * (1.0 + alpha*f_coef[n]))

    IC_prev = np.zeros_like(IC)   # previous IC for criterion A
    SI_prev = np.zeros(K)         # previous scalar irradiance for criterion B

    for it in range(1, max_iter + 1):
        # ── Step 6a: total source at cell centres ─────────────────────────────
        J = Q + S   # shape (N+1, M+1, K)

        # ── Step 6b & 6c: sweeps ─────────────────────────────────────────────
        #
        # The sweep order matches dom_rt.py exactly:
        #   1. Sweep ALL downward directions first.
        #   2. Compute the ground BC from the JUST-COMPLETED downward sweep
        #      (IC_new at k=K+1 for downward directions).
        #   3. Sweep ALL upward directions using that ground BC.
        #
        # This is physically correct: the downward collided flux that hits
        # the ground IN THIS ITERATION is immediately reflected upward.
        # Using the previous iteration's IC for the ground BC (as done
        # naively) causes slower convergence and more iterations are needed.

        IC_new = np.zeros((N + 1, M + 1, K + 2))

        # Step 3 (C4-P4 Slide 7): Sweep downwards.
        # Upper BC: IC = 0 at k=1 (no collided radiation enters from above).
        # IC(Lk+1, Omega_ij) = a_ij * IC(Lk, Omega_ij) - b_ij * J(Lk+1/2, Omega_ij)
        for n in range(1, half + 1):
            for m in range(1, M + 1):
                IC_new[n, m, 1] = 0.0          # upper BC: no collided at top
                for k in range(1, K + 1):
                    IC_new[n, m, k+1] = max(
                        a_coef[n] * IC_new[n, m, k] - b_coef[n] * J[n, m, k-1],
                        0.0)

        # Step 4 (C4-P4 Slide 7): Ground BC for upward sweep.
        # Use the downward IC from THIS sweep (IC_new), not the previous one.
        # This reflects the downward collided flux that arrives at the ground
        # in the current iteration immediately back upward.
        F_col_gnd = 0.0
        for n in range(1, half + 1):
            for m in range(1, M + 1):
                F_col_gnd += w_mu[n] * w_phi * abs(mu[n]) * IC_new[n, m, K+1]
        IC_gnd_refl = (rho_g / math.pi) * F_col_gnd

        # Step 5 (C4-P4 Slide 7): Sweep upwards.
        # IC(Lk, Omega_ij) = c_ij * IC(Lk+1, Omega_ij) + d_ij * J(Lk+1/2, Omega_ij)
        for n in range(half + 1, N + 1):
            for m in range(1, M + 1):
                IC_new[n, m, K+1] = IC_gnd_refl   # lower BC from this sweep
                for k in range(K, 0, -1):
                    IC_new[n, m, k] = max(
                        c_coef[n] * IC_new[n, m, k+1] + d_coef[n] * J[n, m, k-1],
                        0.0)

        IC = IC_new

        # ── Step 6d: new multiple-scattering source ───────────────────────────
        # S[i,j,k] = (1/pi) * SUM_{n=1}^{N} SUM_{m=1}^{M}
        #                      w_n * w_phi * Gamma(nm->ij) * IC_c[n,m,k]
        IC_c = 0.5 * (IC[:, :, 1:K+1] + IC[:, :, 2:K+2])  # cell centres

        S_new = np.zeros((N + 1, M + 1, K))
        for i in range(1, N + 1):
            for j in range(1, M + 1):
                s = np.zeros(K)
                for n in range(1, N + 1):
                    for m in range(1, M + 1):
                        s += w_mu[n] * w_phi * G_qq[n, m, i, j] * IC_c[n, m, :]
                S_new[i, j, :] = s / math.pi

        # ── Step 6e: dual convergence check (identical to dom_rt.py) ────────
        #
        # CRITERION A — max relative change in IC at the BOUNDARY edges
        #   (k=1 and k=K+1) for directions with |mu| >= cos(70 deg).
        #   This ensures the boundary fluxes (Ref and Trans) are converged.
        #
        # CRITERION B — max relative change in the SCALAR IRRADIANCE PROFILE
        #   of the collided field over all K cell centres:
        #   SI[k] = sum_{n,m} w_mu[n] * w_phi * IC_c[n,m,k]   (no |mu|)
        #   max_k |SI_new[k] - SI_prev[k]| / |SI_new[k]|  <  tol
        #   This ensures the INTERIOR radiation density field has converged,
        #   which is required for the scalar irradiance energy balance to
        #   be accurate.  Without criterion B, NIR imbalances reach -40 to -60%.
        #
        # BOTH must be satisfied simultaneously before stopping.

        # ── Step 6e: dual convergence check (identical to dom_rt.py) ────────
        #
        # CRITERION A — max relative change in IC at the BOUNDARY edges
        #   (k=1 and k=K+1) for directions with |mu| >= cos(70 deg) = 0.342.
        #   This is a POINTWISE max, not a sum — it is strict.
        #
        # CRITERION B — max relative change in the collided SCALAR IRRADIANCE
        #   profile over all K cell centres:
        #   SI[k] = sum_{n,m} w_mu[n] * w_phi * IC_c[n,m,k]  (no |mu| weight)
        #   max_k |SI_new[k] - SI_prev[k]| / |SI_new[k]|  <  tol
        #
        # BOTH must pass simultaneously.  Criterion B keeps the solver running
        # until the INTERIOR radiation density field is truly converged — this
        # is essential for the scalar irradiance energy balance to be accurate.

        MU_MIN = math.cos(math.radians(70.0))  # ~ 0.342

        # Criterion A: pointwise max relative change at boundaries
        crit_A = 0.0
        for k_idx in [1, K + 1]:
            for n in range(1, N + 1):
                if abs(mu[n]) < MU_MIN:
                    continue        # skip near-horizontal (poorly resolved)
                for m in range(1, M + 1):
                    denom  = max(abs(IC[n, m, k_idx]), 1e-30)
                    change = abs(IC[n, m, k_idx] - IC_prev[n, m, k_idx]) / denom
                    crit_A = max(crit_A, change)

        # Criterion B: max relative change in scalar irradiance profile
        SI_new = np.zeros(K)
        for n in range(1, N + 1):
            for m in range(1, M + 1):
                SI_new += w_mu[n] * w_phi * IC_c[n, m, :]   # no |mu| weight

        if it == 1:
            crit_B = 1.0            # no previous value on first iteration
        else:
            denom  = np.maximum(np.abs(SI_new), 1e-30)
            crit_B = float(np.max(np.abs(SI_new - SI_prev) / denom))

        SI_prev   = SI_new.copy()
        IC_prev   = IC.copy()
        converged = (crit_A < tol and crit_B < tol)

        if it == 1:
            print(f"    iter  1  (initial sweep, S=0)")
        else:
            print(f"    iter {it:3d}  bnd={crit_A:.2e}  S={crit_B:.2e}"
                  + ("  ✓" if converged else ""))

        S = S_new

        if it > 1 and converged:
            print(f"    Converged in {it} iterations.\n")
            break
    else:
        print(f"    WARNING: did not converge in {max_iter} iterations.\n")

    # ── Compute irradiance profiles from converged IC ─────────────────────────
    Fd_col = np.zeros(K + 2)
    Fu_col = np.zeros(K + 2)
    for k in range(1, K + 2):
        for n in range(1, half + 1):
            for m in range(1, M + 1):
                Fd_col[k] += w_mu[n] * w_phi * abs(mu[n]) * IC[n, m, k]
        for n in range(half + 1, N + 1):
            for m in range(1, M + 1):
                Fu_col[k] += w_mu[n] * w_phi * abs(mu[n]) * IC[n, m, k]

    return IC, it, Fd_col, Fu_col


# =============================================================================
# SECTION 7 — ENERGY BALANCE
# =============================================================================


def compute_scalar_irradiance(I0_dir, I0, IC, mu, w_mu, w_phi, N, M, K):
    """
    Compute the scalar irradiance S(k+1/2) at every cell centre.

    The scalar irradiance is the integral of intensity over ALL directions
    WITHOUT a cosine (|mu|) weight:

        S(k) = integral I(L_k, Omega) dOmega

    In discrete form:
        S_gauss(k) = sum_{n=1}^{N} sum_{m=1}^{M}  w_mu[n] * w_phi * [I0[n,m,k] + IC[n,m,k]]

    The DIRECT BEAM is stored separately as I0_dir (a scalar, not on the Gauss
    grid).  Its contribution to the scalar irradiance is simply the beam
    intensity itself (the delta-function in angle integrates to 1):
        S_direct(k) = I0_dir_c[k]

    Total:  S[k] = S_gauss[k] + S_direct[k]

    This quantity is used in two places:
      1. Convergence criterion B: the solver must run until S itself has
         converged, not just the boundary fluxes.
      2. Leaf absorption:  Abs_leaves = G * (1 - omega_L) * dL * sum_k S[k]

    PARAMETERS
    ----------
    I0_dir : ndarray (K+2,)          direct-beam intensity at cell edges
    I0     : ndarray (N+1,M+1,K+2)   diffuse uncollided intensity at edges
    IC     : ndarray (N+1,M+1,K+2)   collided intensity at cell edges
    mu, w_mu, w_phi, N, M, K : quadrature and grid parameters

    RETURNS
    -------
    S : ndarray (K+2,)   scalar irradiance at cell centres, index 1..K valid
    """
    # Cell-centre values (mean of surrounding edges)
    I0_dir_c = 0.5 * (I0_dir[1:K+1] + I0_dir[2:K+2])            # shape (K,)
    I0_c     = 0.5 * (I0[:, :, 1:K+1] + I0[:, :, 2:K+2])        # (N+1,M+1,K)
    IC_c     = 0.5 * (IC[:, :, 1:K+1] + IC[:, :, 2:K+2])        # (N+1,M+1,K)

    S = np.zeros(K + 2)   # index 0 unused; 1..K valid
    for k in range(1, K + 1):
        # Gauss directions: diffuse uncollided + collided (NO |mu| weight)
        for n in range(1, N + 1):
            for m in range(1, M + 1):
                S[k] += w_mu[n] * w_phi * (I0_c[n, m, k-1] + IC_c[n, m, k-1])
        # Direct beam: beam intensity at cell centre (no |mu| factor)
        S[k] += I0_dir_c[k-1]
    return S

def compute_energy_balance(band_name, LAI,
                            Fd_dir, Fd_dif, Fu_unc,
                            Fd_col, Fu_col,
                            S_total,
                            I0_dir, I0, IC,
                            mu, w_mu, w_phi, N,
                            G, omega_L, rho_g, K, dL, F_in):
    """
    Compute the canopy energy balance using the SCALAR IRRADIANCE method,
    matching dom_rt.py exactly.

    All three quantities are computed INDEPENDENTLY from the solution fields:

        Ref        = Fu_total(k=1) / F_in
                     Upward flux leaving the canopy top.

        Trans      = [Fd_total(K+1) - Fu_total(K+1)] / F_in
                     Net downward flux at the ground (soil absorption).

        Abs_leaves = G * (1 - omega_L) * dL * sum_{k=1}^{K} S(k+1/2) / F_in
                     Leaf absorption via the scalar irradiance integral,
                     where S is computed by compute_scalar_irradiance().

    The energy imbalance  Ref + Trans + Abs_leaves - 1  is a GENUINE diagnostic.
    For RED: ~+0.1%.  For NIR: up to ~-2% depending on LAI and convergence.

    NOTE: This method requires the scalar irradiance S to be FULLY CONVERGED,
    which is enforced by the dual convergence criterion in solve_collided()
    (both boundary fluxes AND S must satisfy the tolerance).  Without dual
    convergence, the NIR imbalance can reach -40% or more.

    PARAMETERS
    ----------
    S_total : ndarray (K+2,)  scalar irradiance from compute_scalar_irradiance()
    (all others: flux arrays and physical parameters)

    RETURNS
    -------
    dict with keys: Ref, Trans, Abs, Abs_leaves, Abs_soil,
                    imb_pct, ref_unc, ref_col, tra_unc, tra_col
    """
    # ── Total irradiance profiles ──────────────────────────────────────────────
    Fd_tot = Fd_dir + Fd_dif + Fd_col
    Fu_tot = Fu_unc + Fu_col

    # ── Reflectance — INDEPENDENT ─────────────────────────────────────────────
    Ref = Fu_tot[1] / F_in

    # ── Transmittance — INDEPENDENT ───────────────────────────────────────────
    # Net downward flux crossing the soil surface:
    #   Trans = Fd_tot(K+1) - Fu_tot(K+1)
    # where Fu_tot(K+1) = rho_g * Fd_tot(K+1) (Lambertian ground).
    # So Trans = (1-rho_g)*Fd_tot(K+1) = soil absorption.
    Trans = (Fd_tot[K+1] - Fu_tot[K+1]) / F_in

    # ── Soil absorption — INDEPENDENT ─────────────────────────────────────────
    # The soil absorbs (1-rho_g) of the total downward flux reaching it.
    # Computed directly: Abs_soil = (1-rho_g)*Fd_tot[K+1]/F_in.
    # This equals Trans to within the GL quadrature accuracy (~0.01%).
    Abs_soil = (1.0 - rho_g) * Fd_tot[K+1] / F_in

    # ── Leaf absorption via scalar irradiance — INDEPENDENT ───────────────────
    # S(k) = integral of total intensity over all angles (no cosine weight).
    # Physical absorption rate per unit LAI = G * (1-omega_L) * S(L).
    S_vals     = S_total[1:K+1]   # cell centres k=1..K
    Abs_leaves = G * (1.0 - omega_L) * dL * np.sum(S_vals) / F_in

    # ── Total absorption ───────────────────────────────────────────────────────
    Abs = Abs_leaves + Abs_soil

    # ── Energy imbalance — GENUINE DIAGNOSTIC ─────────────────────────────────
    # The correct balance is:  Ref + Trans + Abs_leaves = 1.0
    # (Trans = Abs_soil, so do NOT add both — that would double-count the soil.)
    imb     = Ref + Trans + Abs_leaves - 1.0
    imb_pct = imb * 100.0

    # ── Component breakdown: uncollided vs collided ───────────────────────────
    #
    # Each quantity is split into its uncollided (I0) and collided (IC) parts.
    # This matches the column structure of the output Word table.

    # Reflectance components
    ref_unc = Fu_unc[1] / F_in
    ref_col = Fu_col[1] / F_in

    # Transmittance components (net flux at ground, per field)
    tra_unc = (Fd_dir[K+1] + Fd_dif[K+1] - Fu_unc[K+1]) / F_in
    tra_col = (Fd_col[K+1]               - Fu_col[K+1])  / F_in

    # Soil absorption components ((1-rho_g) * downward flux, per field)
    soil_unc = (1.0 - rho_g) * (Fd_dir[K+1] + Fd_dif[K+1]) / F_in
    soil_col = (1.0 - rho_g) *  Fd_col[K+1]                 / F_in

    # Leaf absorption components (scalar irradiance, split by uncollided / collided)
    # Cell-centre intensities
    I0_dir_c = 0.5 * (I0_dir[1:K+1] + I0_dir[2:K+2])          # shape (K,)
    I0_c     = 0.5 * (I0[:, :, 1:K+1] + I0[:, :, 2:K+2])      # (N+1,M+1,K)
    IC_c     = 0.5 * (IC[:, :, 1:K+1] + IC[:, :, 2:K+2])      # (N+1,M+1,K)
    # Scalar irradiance from uncollided field (Gauss directions + direct beam)
    SI_unc = np.zeros(K)
    SI_col = np.zeros(K)
    # Determine M from w_phi
    # Number of azimuthal directions inferred from w_phi = 2*pi/M
    M_eb = int(round(2.0 * math.pi / w_phi))
    for n in range(1, N + 1):
        for m in range(1, M_eb + 1):
            SI_unc += w_mu[n] * w_phi * I0_c[n, m, :]
            SI_col += w_mu[n] * w_phi * IC_c[n, m, :]
    SI_unc += I0_dir_c   # direct-beam contribution to uncollided scalar irradiance
    leaf_unc = G * (1.0 - omega_L) * dL * np.sum(SI_unc) / F_in
    leaf_col = G * (1.0 - omega_L) * dL * np.sum(SI_col) / F_in

    return dict(
        Ref        = Ref,
        Trans      = Trans,
        Abs        = Abs,
        Abs_leaves = Abs_leaves,
        Abs_soil   = Abs_soil,
        imb_pct    = imb_pct,
        # Reflectance split
        ref_unc    = ref_unc,    ref_col  = ref_col,
        # Transmittance split
        tra_unc    = tra_unc,    tra_col  = tra_col,
        # Soil absorption split
        soil_unc   = soil_unc,   soil_col = soil_col,
        # Leaf absorption split
        leaf_unc   = leaf_unc,   leaf_col = leaf_col,
    )





# =============================================================================
# SECTION 8 — DOM INTERPOLATION INTO ARBITRARY VIEW DIRECTION
# =============================================================================

def dom_interpolate_view(
        theta_v_deg, phi_v_deg,
        mu, phi_q, w_mu, w_phi, N, M, K,
        G, dL, rho_g, omega_L, tau_L,
        solar_mu, solar_phi,
        I0_dir, I0, IC, F_in):
    """
    Interpolate total intensity I(L=0, Omega_V) into an arbitrary upward
    view direction Omega_V = (theta_v, phi_v) using the DOM method.

    This implements exactly the 9-step procedure described in the companion
    document (and slides 17–22 of the lecture notes):

    Step 1 : Evaluate cross-sections G(Omega_V), Gamma(solar->Omega_V),
             Gamma(nm->Omega_V) for all Gauss pairs (n,m)
    Step 2 : Uncollided ground BC in Omega_V
    Step 3 : Beer-Lambert upward sweep for uncollided intensity
    Step 4 : First Collision Source Q1+Q2+Q3 at cell centres in Omega_V
    Step 5 : Multiple-scattering source S at cell centres in Omega_V
             (using converged IC from Gauss solve)
    Step 6 : Total source J = Q + S
    Step 7 : Collided ground BC in Omega_V
    Step 8 : Single upward Diamond-Difference sweep -> IC(L=0, Omega_V)
    Step 9 : Add uncollided and collided: I_total = I0 + IC

    PARAMETERS
    ----------
    theta_v_deg : float   view polar angle from nadir (degrees, 0..90)
    phi_v_deg   : float   view azimuth angle (degrees, 0..360)
    (all other parameters: pre-computed fields from the Gauss solve)

    RETURNS
    -------
    float : total intensity I(L=0, Omega_V)
            Divide by (F_in / pi) to get BRF or HDRF.
    """
    # Convert to radians; mu_v > 0 (upward)
    mu_v  = math.cos(math.radians(theta_v_deg))
    phi_v = math.radians(phi_v_deg)

    half = N // 2

    # ── Step 1: Gamma tables for this view direction ───────────────────────────
    # Gamma(solar -> Omega_V)
    Gam_sol_v = gamma_scalar(solar_mu, solar_phi, mu_v, phi_v, omega_L, tau_L)

    # Gamma(nm -> Omega_V) for all Gauss directions: vectorised using numpy
    mu_arr  = mu[1:N+1]           # (N,)
    phi_arr = phi_q[1:M+1]        # (M,)
    sin_n   = np.sqrt(np.maximum(0.0, 1.0 - mu_arr**2))    # (N,)
    sin_v   = math.sqrt(max(0.0, 1.0 - mu_v**2))

    # cos(beta) for all (n,m) pairs: shape (N,M)
    cos_beta = np.clip(
        mu_arr[:, np.newaxis] * mu_v
        + sin_n[:, np.newaxis] * sin_v * np.cos(phi_arr[np.newaxis, :] - phi_v),
        -1.0, 1.0
    )
    beta     = np.arccos(cos_beta)
    Gam_nm_v = ((omega_L / (3.0 * math.pi)) * (np.sin(beta) - beta * cos_beta)
                + (tau_L / 3.0) * cos_beta)   # (N, M)

    # ── Step 2: Uncollided ground BC in Omega_V ────────────────────────────────
    F_dn_gnd = 0.0
    for n in range(1, half + 1):
        for m in range(1, M + 1):
            F_dn_gnd += w_mu[n] * w_phi * abs(mu[n]) * I0[n, m, K+1]
    F_dn_gnd += abs(solar_mu) * I0_dir[K+1]
    I0_gnd_v = (rho_g / math.pi) * F_dn_gnd

    # ── Step 3: Uncollided intensity at canopy top in Omega_V ──────────────────
    I0_top_v = I0_gnd_v * math.exp(-G / mu_v * (K * dL))

    # ── Step 4: First Collision Source in Omega_V ──────────────────────────────
    I0_dir_c = 0.5 * (I0_dir[1:K+1] + I0_dir[2:K+2])           # (K,)
    I0_c     = 0.5 * (I0[:, :, 1:K+1] + I0[:, :, 2:K+2])       # (N+1,M+1,K)

    # Q1
    Q1_v = (1.0 / math.pi) * Gam_sol_v * I0_dir_c               # (K,)

    # Q2 (downward diffuse, n=1..N/2) and Q3 (upward, n=N/2+1..N)
    # using einsum for speed: Gam_nm_v has shape (N,M), wI0 has shape (N,M,K)
    wI0_dn  = (w_mu[1:half+1, np.newaxis, np.newaxis] * w_phi
               * I0_c[1:half+1, 1:M+1, :])                       # (N/2,M,K)
    wI0_up  = (w_mu[half+1:N+1, np.newaxis, np.newaxis] * w_phi
               * I0_c[half+1:N+1, 1:M+1, :])                     # (N/2,M,K)

    Q2_v = np.einsum('nm,nmk->k', Gam_nm_v[:half,  :], wI0_dn) / math.pi
    Q3_v = np.einsum('nm,nmk->k', Gam_nm_v[half:,  :], wI0_up) / math.pi

    # ── Step 5: Multiple-scattering source in Omega_V ─────────────────────────
    IC_c    = 0.5 * (IC[:, :, 1:K+1] + IC[:, :, 2:K+2])         # (N+1,M+1,K)
    wIC_all = (w_mu[1:N+1, np.newaxis, np.newaxis] * w_phi
               * IC_c[1:N+1, 1:M+1, :])                           # (N,M,K)
    S_v = np.einsum('nm,nmk->k', Gam_nm_v, wIC_all) / math.pi    # (K,)

    # ── Step 6: Total source ───────────────────────────────────────────────────
    J_v = Q1_v + Q2_v + Q3_v + S_v   # (K,)

    # ── Step 7: Collided ground BC in Omega_V ─────────────────────────────────
    F_col_gnd = 0.0
    for n in range(1, half + 1):
        for m in range(1, M + 1):
            F_col_gnd += w_mu[n] * w_phi * abs(mu[n]) * IC[n, m, K+1]
    IC_gnd_v = (rho_g / math.pi) * F_col_gnd

    # ── Step 8: Single upward DD sweep ────────────────────────────────────────
    half_f_v = 0.5 * G * dL / mu_v
    c_v  = (1.0 - half_f_v) / (1.0 + half_f_v)
    dd_v = (dL / mu_v)       / (1.0 + half_f_v)

    IC_edge_v = np.zeros(K + 2)
    IC_edge_v[K+1] = IC_gnd_v
    for k in range(K, 0, -1):
        IC_edge_v[k] = max(c_v * IC_edge_v[k+1] + dd_v * J_v[k-1], 0.0)

    # ── Step 9: Total intensity at canopy top ──────────────────────────────────
    return I0_top_v + IC_edge_v[1]


# =============================================================================
# SECTION 9 — WRITE ENERGY-BALANCE WORD DOCUMENT
# =============================================================================

def _set_cell_bg(cell, hex_color):
    """Apply a background fill to a table cell (python-docx helper)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def _bold_run(para, text, size_pt=10, color_hex=None):
    """Add a bold run to a paragraph."""
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def write_energy_balance_docx(results_all, output_path):
    """
    Write the energy balance table to a Word document.

    For each band and each LAI value, a table is produced showing:
      - Uncollided reflectance, transmittance
      - Collided reflectance, transmittance
      - Total reflectance, transmittance, leaf absorption, soil absorption
      - Energy balance imbalance (%)

    PARAMETERS
    ----------
    results_all : dict  band_name -> LAI -> energy-balance dict
    output_path : str   path to write the .docx file
    """
    doc = DocxDocument()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    # Title
    title = doc.add_heading("Energy Balance — Discrete Ordinates Method", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    import datetime
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    HEADER_COLOR = "2E75B6"   # blue
    ROW_ALT      = "EBF3FB"   # light blue for alternating rows
    WHITE        = "FFFFFF"

    for band_name, lai_dict in results_all.items():
        doc.add_heading(f"Band: {band_name}", level=2)

        for LAI, eb in lai_dict.items():
            doc.add_heading(f"  LAI = {LAI}", level=3)

            # Table: 3 columns — Quantity | Uncollided | Collided | Total
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            hdr_cells = tbl.rows[0].cells
            for cell, text in zip(hdr_cells,
                                   ["Quantity", "Uncollided", "Collided", "Total"]):
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _bold_run(p, text, size_pt=10, color_hex=WHITE)
                _set_cell_bg(cell, HEADER_COLOR)

            # Data rows — four columns: Quantity | Uncollided | Collided | Total
            # Energy balance: Ref + Trans + Abs_leaves = 1.0
            # Trans = Abs_soil (Lambertian ground), so Abs_soil is shown for
            # physical completeness but NOT added to SUM (would double-count).
            sum_unc = eb['ref_unc'] + eb['tra_unc'] + eb['leaf_unc']
            sum_col = eb['ref_col'] + eb['tra_col'] + eb['leaf_col']
            sum_tot = eb['Ref']     + eb['Trans']   + eb['Abs_leaves']
            rows_data = [
                ("Reflectance",
                 f"{eb['ref_unc']:.5f}",
                 f"{eb['ref_col']:.5f}",
                 f"{eb['Ref']:.5f}"),
                ("Transmittance (net)",
                 f"{eb['tra_unc']:.5f}",
                 f"{eb['tra_col']:.5f}",
                 f"{eb['Trans']:.5f}"),
                ("Leaf Absorption",
                 f"{eb['leaf_unc']:.5f}",
                 f"{eb['leaf_col']:.5f}",
                 f"{eb['Abs_leaves']:.5f}"),
                ("Soil Absorption (=Trans)",
                 f"{eb['soil_unc']:.5f}",
                 f"{eb['soil_col']:.5f}",
                 f"{eb['Abs_soil']:.5f}"),
                ("SUM  Ref+Trans+Leaf",
                 f"{sum_unc:.5f}",
                 f"{sum_col:.5f}",
                 f"{sum_tot:.5f}"),
                ("Energy Imbalance (%)",
                 "", "",
                 f"{eb['imb_pct']:+.4f} %"),
            ]

            for row_idx, (q, unc, col, tot) in enumerate(rows_data):
                row = tbl.add_row()
                bg = ROW_ALT if row_idx % 2 == 0 else WHITE
                for cell, val in zip(row.cells, [q, unc, col, tot]):
                    cell.text = val
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_cell_bg(cell, bg)
                # Bold the label
                row.cells[0].paragraphs[0].runs[0].bold = True

            doc.add_paragraph("")   # spacing after table

    doc.save(output_path)
    print(f"  Energy balance written to: {output_path}")


# =============================================================================
# SECTION 10 — WRITE BRF / HDRF WORD DOCUMENT
# =============================================================================

def write_brf_hdrf_docx(brf_grids, theta_v_arr, phi_v_arr, results_all,
                         f_dir_per_band, output_path):
    """
    Write the BRF / HDRF table to a Word document.

    One table per (band, LAI) combination.  Rows = view polar angles,
    columns = view azimuth angles.

    PARAMETERS
    ----------
    brf_grids    : dict  band_name -> LAI -> ndarray(n_theta, n_phi)
    theta_v_arr  : ndarray  view polar angles from nadir (degrees)
    phi_v_arr    : ndarray  view azimuth angles (degrees)
    results_all  : dict     (for header info)
    f_dir_per_band : dict   band_name -> f_dir
    output_path  : str
    """
    doc = DocxDocument()

    section = doc.sections[0]
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.5)
    section.right_margin  = Inches(0.5)
    from docx.enum.section import WD_ORIENT
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Inches(11)
    section.page_height = Inches(8.5)

    import datetime
    title = doc.add_heading("BRF / HDRF Table — Discrete Ordinates Method", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(
        f"Polar step: {theta_v_arr[1]-theta_v_arr[0]:.0f} deg    "
        f"Azimuth step: {phi_v_arr[1]-phi_v_arr[0]:.0f} deg    "
        f"Directions: {len(theta_v_arr)} x {len(phi_v_arr)} = "
        f"{len(theta_v_arr)*len(phi_v_arr)}")
    doc.add_paragraph("")

    HEADER_COLOR = "2E75B6"
    ROW_COLORS   = ["FFFFFF", "EBF3FB"]

    for band_name, lai_dict in brf_grids.items():
        f_dir = f_dir_per_band[band_name]
        label = "BRF" if abs(f_dir - 1.0) < 1e-6 else "HDRF"

        for LAI, grid in lai_dict.items():
            doc.add_heading(f"{label} — {band_name}  LAI = {LAI}", level=2)

            n_phi = len(phi_v_arr)
            n_theta = len(theta_v_arr)

            # Table: first column = theta label, then one column per phi
            tbl = doc.add_table(rows=1, cols=1 + n_phi)
            tbl.style = 'Table Grid'

            # Header row
            hdr = tbl.rows[0].cells
            hdr[0].text = ""
            p0 = hdr[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _bold_run(p0, f"theta\\phi", size_pt=8, color_hex="FFFFFF")
            _set_cell_bg(hdr[0], HEADER_COLOR)

            for j, phi_val in enumerate(phi_v_arr):
                hdr[j+1].text = ""
                p = hdr[j+1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _bold_run(p, f"{phi_val:.0f}", size_pt=8, color_hex="FFFFFF")
                _set_cell_bg(hdr[j+1], HEADER_COLOR)

            # Data rows
            for i, theta_val in enumerate(theta_v_arr):
                row = tbl.add_row()
                bg = ROW_COLORS[i % 2]

                # Row label
                row.cells[0].text = ""
                p = row.cells[0].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _bold_run(p, f"{theta_val:.0f}", size_pt=8)
                _set_cell_bg(row.cells[0], "D9E1F2")

                for j in range(n_phi):
                    val = grid[i, j]
                    row.cells[j+1].text = f"{val:.4f}"
                    row.cells[j+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in row.cells[j+1].paragraphs[0].runs:
                        run.font.size = Pt(8)
                    _set_cell_bg(row.cells[j+1], bg)

            doc.add_paragraph("")

    doc.save(output_path)
    print(f"  BRF/HDRF table written to: {output_path}")


# =============================================================================
# SECTION 11 — MASTER SOLVE ROUTINE
# =============================================================================

def solve_one_case(band_name, band_params, LAI,
                   theta_sun_deg, phi_sun_deg, f_dir, F_in,
                   N, M, K, tol, max_iter,
                   G_qq=None, G_sol=None, mu=None, phi=None,
                   w_mu=None, w_phi=None):
    """
    Run the complete DOM solver for one (band, LAI) pair.

    Steps performed:
      1.  Quadrature setup (or use cached if supplied)
      2.  Uncollided field
      3.  First Collision Source
      4.  Collided field
      5.  Energy balance

    PARAMETERS
    ----------
    band_name    : str
    band_params  : dict  with keys rho_L, tau_L, rho_g
    LAI          : float
    theta_sun_deg, phi_sun_deg : float  illumination geometry
    f_dir        : float   direct fraction
    F_in         : float   incoming irradiance
    N, M, K      : int     quadrature and layer resolution
    tol, max_iter : float, int   solver settings
    G_qq         : precomputed Gamma table (optional, re-used across LAI)
    G_sol        : precomputed solar-Gauss Gamma table (optional)
    mu, phi, w_mu, w_phi : precomputed quadrature (optional)

    RETURNS
    -------
    dict with all intermediate and final arrays needed for output and
    interpolation.
    """
    rho_L   = band_params["rho_L"]
    tau_L   = band_params["tau_L"]
    rho_g   = band_params["rho_g"]
    omega_L = rho_L + tau_L
    G       = G_extinction()
    dL      = LAI / K

    solar_mu  = math.cos(math.radians(theta_sun_deg))
    solar_phi = math.radians(phi_sun_deg)

    t0 = time.time()

    # Quadrature
    if mu is None:
        mu, phi, w_mu, w_phi = setup_quadrature(N, M)
    else:
        pass

    # Gamma tables
    if G_qq is None:
        G_qq = precompute_gamma_quad(mu, phi, N, M, omega_L, tau_L)
    if G_sol is None:
        G_sol = precompute_gamma_solar(solar_mu, solar_phi, mu, phi, N, M,
                                        omega_L, tau_L)

    print(f"  [{band_name}  LAI={LAI}] Uncollided sweep ...")
    I0_dir, I0, Fd_dir, Fd_dif, Fu_unc = solve_uncollided(
        mu, phi, w_mu, w_phi, N, M, K,
        solar_mu, solar_phi, f_dir, F_in, rho_g, G, dL)

    print(f"  [{band_name}  LAI={LAI}] First Collision Source ...")
    Q = compute_first_collision_source(
        mu, phi, w_mu, w_phi, N, M, K,
        solar_mu, solar_phi, I0_dir, I0,
        G_sol, G_qq, omega_L)

    print(f"  [{band_name}  LAI={LAI}] Collided solver ...")
    IC, n_iter, Fd_col, Fu_col = solve_collided(
        mu, phi, w_mu, w_phi, N, M, K,
        G, dL, rho_g, Q, G_qq, omega_L, tol, max_iter)

    # Compute scalar irradiance (needed for energy balance and was tracked
    # inside solve_collided for convergence; recompute here for clarity)
    S_total = compute_scalar_irradiance(
        I0_dir, I0, IC, mu, w_mu, w_phi, N, M, K)

    eb = compute_energy_balance(
        band_name, LAI,
        Fd_dir, Fd_dif, Fu_unc,
        Fd_col, Fu_col,
        S_total,
        I0_dir, I0, IC,
        mu, w_mu, w_phi, N,
        G, omega_L, rho_g, K, dL, F_in)

    elapsed = time.time() - t0
    print(f"  Done in {elapsed:.1f}s  |  {n_iter} iterations  "
          f"|  imbalance = {eb['imb_pct']:+.3f}%\n")

    return dict(
        eb=eb, IC=IC, I0_dir=I0_dir, I0=I0,
        mu=mu, phi=phi, w_mu=w_mu, w_phi=w_phi,
        G_qq=G_qq, G_sol=G_sol,
        G=G, omega_L=omega_L, tau_L=tau_L, rho_g=rho_g,
        solar_mu=solar_mu, solar_phi=solar_phi,
        K=K, dL=dL, F_in=F_in, n_iter=n_iter,
    )


# =============================================================================
# SECTION 12 — MAIN ENTRY POINT
# =============================================================================

def main():
    print("=" * 70)
    print("  DOM Radiative Transfer Solver  —  Student Edition")
    print("=" * 70)
    print()

    # ── Ask user whether to do interpolation ──────────────────────────────────
    while True:
        ans = input(
            "Perform BRF/HDRF interpolation into the full angular grid? "
            "(1 = yes, 0 = no): "
        ).strip()
        if ans in ("0", "1"):
            do_interp = (ans == "1")
            break
        print("  Please enter 0 or 1.")
    print()

    # ── Build view-direction grid (used only if do_interp) ────────────────────
    theta_v_arr = np.arange(THETA_V_STEP, 90, THETA_V_STEP, dtype=float)
    phi_v_arr   = np.arange(0, 360, PHI_V_STEP, dtype=float)
    n_dirs = len(theta_v_arr) * len(phi_v_arr)

    if do_interp:
        print(f"  Interpolation grid: {len(theta_v_arr)} theta x "
              f"{len(phi_v_arr)} phi = {n_dirs} directions")
        print()

    # ── Main loop over bands and LAI ──────────────────────────────────────────
    results_all = {}   # band -> LAI -> eb dict
    brf_grids   = {}   # band -> LAI -> 2D array  (only if do_interp)

    for band_name, bparams in SPECTRAL_BANDS.items():
        f_dir   = bparams["f_dir"]
        omega_L = bparams["rho_L"] + bparams["tau_L"]
        tau_L   = bparams["tau_L"]
        rho_g   = bparams["rho_g"]

        print(f"{'='*70}")
        print(f"  BAND: {band_name}   omega_L={omega_L:.4f}   "
              f"rho_g={rho_g}   f_dir={f_dir}")
        print(f"{'='*70}\n")

        results_all[band_name] = {}
        if do_interp:
            brf_grids[band_name] = {}

        # Pre-compute quadrature once per band
        mu, phi, w_mu, w_phi = setup_quadrature(N_MU, N_PHI)

        # Pre-compute Gamma tables once per band (shared across all LAI)
        print(f"  Pre-computing Gamma tables for {band_name} ...")
        G_qq = precompute_gamma_quad(mu, phi, N_MU, N_PHI, omega_L, tau_L)
        solar_mu_tmp  = math.cos(math.radians(THETA_SUN_DEG))
        solar_phi_tmp = math.radians(PHI_SUN_DEG)
        G_sol = precompute_gamma_solar(
            solar_mu_tmp, solar_phi_tmp,
            mu, phi, N_MU, N_PHI, omega_L, tau_L)
        print()

        for LAI in LAI_VALUES:
            print(f"  ── LAI = {LAI} ──")
            d = solve_one_case(
                band_name, bparams, LAI,
                THETA_SUN_DEG, PHI_SUN_DEG, f_dir, F_IN,
                N_MU, N_PHI, N_LAY, TOL, MAX_ITER,
                G_qq=G_qq, G_sol=G_sol, mu=mu, phi=phi,
                w_mu=w_mu, w_phi=w_phi)

            results_all[band_name][LAI] = d["eb"]

            # ── Optional: interpolation ────────────────────────────────────────
            if do_interp:
                label = "BRF" if abs(f_dir - 1.0) < 1e-6 else "HDRF"
                print(f"  Interpolating {label} for {band_name} LAI={LAI} "
                      f"({n_dirs} directions) ...")
                brf_norm = F_IN / math.pi
                grid = np.zeros((len(theta_v_arr), len(phi_v_arr)))
                t_interp = time.time()

                for i, th in enumerate(theta_v_arr):
                    for j, ph in enumerate(phi_v_arr):
                        I_total = dom_interpolate_view(
                            th, ph,
                            mu, phi, w_mu, w_phi, N_MU, N_PHI, N_LAY,
                            d["G"], d["dL"], rho_g, omega_L, tau_L,
                            d["solar_mu"], d["solar_phi"],
                            d["I0_dir"], d["I0"], d["IC"], F_IN)
                        grid[i, j] = I_total / brf_norm

                print(f"  Done in {time.time()-t_interp:.1f}s\n")
                brf_grids[band_name][LAI] = grid

    # ── Write output documents ────────────────────────────────────────────────
    print("=" * 70)
    print("  Writing output documents ...")
    print("=" * 70)

    write_energy_balance_docx(results_all, ENERGY_BALANCE_DOCX)

    if do_interp:
        f_dir_per_band = {b: p["f_dir"] for b, p in SPECTRAL_BANDS.items()}
        write_brf_hdrf_docx(
            brf_grids, theta_v_arr, phi_v_arr,
            results_all, f_dir_per_band, BRF_HDRF_DOCX)

    print()
    print("=" * 70)
    print("  Complete.")
    if do_interp:
        print(f"  -> {ENERGY_BALANCE_DOCX}")
        print(f"  -> {BRF_HDRF_DOCX}")
    else:
        print(f"  -> {ENERGY_BALANCE_DOCX}")
    print("=" * 70)


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    main()
